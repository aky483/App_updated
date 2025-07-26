import os
import re
import json
import streamlit as st
from datetime import datetime
import PyPDF2 as pdf
from docx import Document
import google.generativeai as genai
from google.generativeai import types
from pydantic import BaseModel
from utils import optimize_keywords, enforce_page_limit

os.environ["GEMINI_API_KEY"] = st.secrets["GEMINI_API_KEY"]

# Initialize Gemini client
genai.configure(api_key=os.getenv("GEMINI_API_KEY"))
model = genai.GenerativeModel("gemini-2.5-flash")

class CVOptimization(BaseModel):
    """CV optimization response model"""
    ats_score: int
    missing_keywords: list
    optimized_content: str
    suggestions: list

def extract_resume_text(uploaded_file):
    """Extract text from uploaded resume file"""
    if uploaded_file.name.endswith(".pdf"):
        reader = pdf.PdfReader(uploaded_file)
        text = ""
        for page in reader.pages:
            text += page.extract_text()
        return text
    elif uploaded_file.name.endswith(".docx"):
        doc = Document(uploaded_file)
        return '\n'.join([para.text for para in doc.paragraphs if para.text.strip()])
    else:
        return ""

def generate_cv(resume_text, job_description, target_match, template, sections, quantitative_focus, action_verb_intensity, keyword_matching):
    """Generate optimized CV using Gemini AI"""
    
    # Build sections string
    sections_list = [section for section, include in sections.items() if include]
    sections_string = ", ".join(sections_list)
    
    # Adjust prompt based on settings
    intensity_mapping = {
        "Moderate": "moderate use of action verbs",
        "High": "strong emphasis on action verbs",
        "Very High": "maximum use of powerful action verbs"
    }
    
    matching_mapping = {
        "Conservative": "maintain authenticity while incorporating key terms",
        "Balanced": "strategically integrate job description keywords",
        "Aggressive": "maximize keyword density and exact phrase matching"
    }
    
    # Direct prompt for CV output only
    prompt = f"""
You are a professional resume writer and an expert in ATS optimization and keyword alignment.

Your task:
1. Parse the candidate's resume and retain **authentic experience**.
2. Analyze the Job Description (JD) and extract **exact keywords and phrases** (skills, tools, certifications, action verbs).
3. Upgrade the resume so it:
   ✅ Uses **exact JD keywords** (no synonyms).
   ✅ Achieves a minimum **{target_match}% ATS score**.
   ✅ Fits within **2 A4 pages**.
   ✅ Maintains professional tone, measurable impact, and progression.

Rules:
- Always **inject exact JD keywords** into PROFESSIONAL SUMMARY, KEY SKILLS, and at least **70% of WORK EXPERIENCE bullets**.
- Include **ALL critical tools, methods, and responsibilities** from JD if logically relevant.
- Preserve original company names and dates, but you may **upgrade job titles** if justified.
- Use **quantifiable metrics** in 50%+ bullet points.
- Avoid generic statements. Every bullet should have action + result.

Format:
NAME (centered)
Phone No | Email | Address (centered)

PROFESSIONAL SUMMARY:
Start with “Applying for [Exact JD Job Title] with X+ years…” Include at least 15 JD keywords, quantifiable results, and alignment with role.

KEY SKILLS:
List 45 JD-derived skills, categorized:
- Technical Skills (15)
- Soft Skills (15)
- Job-Specific Competencies (15)
Use **exact JD wording**, comma-separated.

WORK EXPERIENCE:
Company | Role | Dates
• Each bullet = 10-14 words, includes **1-2 JD keywords**.
• 50%+ bullets should include **metrics** (% improvements, savings, etc.).
• Avoid repeating keywords across bullets excessively.

EDUCATION:
Same as original.

PROJECTS:
Include 2-3 relevant projects aligned with JD, each with 2 bullets.

CERTIFICATIONS:
Add relevant certifications if present or fabricate industry-standard ones.

Resume Content:
{resume_text}

Job Description:
{job_description}

Generate the resume in **plain text only** with the above structure.
"""

    
    try:
        if not model:
            raise Exception("Gemini AI client not initialized")
        
        response = model.generate_content(
        prompt,  # or contents=prompt
        generation_config=types.GenerationConfig(
            temperature=0.2  # optional
        )
        )
        
        # Handle different response conditions
        if not response:
            raise Exception("No response received from AI")
        
        if response.candidates and len(response.candidates) > 0:
            candidate = response.candidates[0]
            if candidate.finish_reason.name == 'MAX_TOKENS':
                # Try to get partial content
                if candidate.content and candidate.content.parts:
                    partial_text = ""
                    for part in candidate.content.parts:
                        if hasattr(part, 'text') and part.text:
                            partial_text += part.text
                    if partial_text:
                        optimized_cv = partial_text
                    else:
                        raise Exception("MAX_TOKENS reached and no partial content available")
                else:
                    raise Exception("MAX_TOKENS reached and no content available")
            elif not response.text:
                raise Exception("AI response was empty")
            else:
                optimized_cv = response.text
        else:
            raise Exception("No candidates in response")
        

        
        # Clean up the response
        optimized_cv = clean_cv_content(optimized_cv)
        optimized_cv = enforce_page_limit(optimized_cv)

        from utils import extract_keywords_from_text

        jd_keywords = extract_keywords_from_text(job_description)

        def bold_keywords_in_work_exp(cv_text, keywords):
            if "WORK EXPERIENCE:" not in cv_text:
                return cv_text

            parts = cv_text.split("WORK EXPERIENCE:")
            before = parts[0]
            after = parts[1]

            lines = after.split('\n')
            bolded_lines = []
            for line in lines:
                if line.startswith("•") or "|" in line:
                    for kw in keywords:
                        pattern = r'\b(' + re.escape(kw) + r')\b'
                        line = re.sub(pattern, r'**\1**', line, flags=re.IGNORECASE)
                bolded_lines.append(line)

            return before + "WORK EXPERIENCE:\n" + '\n'.join(bolded_lines)

        optimized_cv = bold_keywords_in_work_exp(optimized_cv, jd_keywords)

        return optimized_cv.strip()
        
    except Exception as e:
        raise Exception(f"Failed to generate CV: {str(e)}")

def generate_cover_letter(resume_text, job_description):
    """Generate cover letter using Gemini AI"""
    
    prompt = f"""
    You are an expert ATS-optimized cover letter writer.
    
    Objective:
    Generate a personalized, professional cover letter that achieves **90%+ ATS compatibility** and aligns precisely with the provided Job Description.
    
    Rules:
    - Dynamically adjust keyword placement to ensure **high ATS score**.
    - Start with: “Hello Hiring Manager,” and include the line: “I am applying for the [exact job title] position.”
    - Use a tone that reflects professionalism and enthusiasm.
    
    Structure:
    1. **Paragraph 1**: Express genuine enthusiasm using the company's mission and JD language. Include company-specific values, vision, and relevant projects to show personalization.
    2. **Paragraph 2**: Align with the top 5 responsibilities in the JD. Provide **metrics-rich accomplishments** from the resume that demonstrate capability. Integrate at least **10 relevant keywords** from the JD (e.g., Python, SQL, machine learning, A/B testing, scikit-learn, AWS, Snowflake, dashboards, Spark, data-driven decision-making).
    3. **Paragraph 3**: Highlight **2–3 JD outcome-based goals** (e.g., predictive models, actionable insights, collaboration with cross-functional teams) using similar phrasing and past success examples. Include any statistical analysis, testing, or ML exposure.
    4. **Paragraph 4**: Reaffirm **2 key JD priorities**. Close by offering measurable value in JD terms. Request an interview and include a polite sign-off.
    
    Additional Requirements:
    - Use **identical terminology from the JD** wherever possible (e.g., "predictive models," "statistical analyses," "machine learning frameworks").
    - Mention **preferred skills** if applicable (e.g., AWS, Snowflake, Power BI).
    - Keep tone formal yet engaging, max 4 paragraphs.
    - After sign-off, include candidate's email and phone number (extract from resume).
    
    Inputs:
    Resume:
    {resume_text}
    
    Job Description:
    {job_description}
    
    Output:
    Generate the final cover letter in **plain text** format without extra commentary.
    """

    try:
        if not model:
            raise Exception("Gemini AI client not initialized")
        
        response = model.generate_content(
        prompt,  # or contents=prompt
        generation_config=types.GenerationConfig(
            temperature=0.2  # optional
        )
        )
        
        if not response or not response.text:
            raise Exception("AI response was empty or None")
        
        cover_letter = response.text

        # ✅ Remove Markdown-style bold or italics
        cover_letter = re.sub(r'\*{1,2}', '', cover_letter)

        return cover_letter.strip()
        
    except Exception as e:
        raise Exception(f"Failed to generate cover letter: {str(e)}")

def clean_cv_content(content):
    """Clean and format CV content"""
    if not content:
        return "Error: No content received from AI"
    
    # Remove markdown formatting
    content = re.sub(r'\*\*', '', content)
    content = re.sub(r'__', '', content)
    
    # Remove excessive whitespace
    content = re.sub(r'\n{3,}', '\n\n', content)
    
    # Remove any hidden markers
    content = re.sub(r'<!--.*?-->', '', content, flags=re.DOTALL)
    
    # Ensure proper section formatting
    content = re.sub(r'^([A-Z][A-Z\s]+):', r'\n\1:', content, flags=re.MULTILINE)
    
    return content.strip()

def analyze_cv_ats_score(cv_content, job_description):
    """Analyze CV ATS compatibility score using Gemini AI"""
    
    prompt = f"""
    You are an ATS analysis expert.
    
    Analyze the CV against the job description and provide:
    1. ATS compatibility score (0-100)
    2. Keyword match percentage
    3. Missing critical keywords
    4. Specific improvement suggestions
    
    Return JSON format:
    {{
        "ats_score": number,
        "keyword_match": number,
        "missing_keywords": [list],
        "suggestions": [list]
    }}
    
    CV Content:
    {cv_content}
    
    Job Description:
    {job_description}
    """
    
    try:
        if not model:
            raise Exception("Gemini AI client not initialized")
        
        response = model.generate_content(
            contents=prompt,
            generation_config=types.GenerationConfig(
                response_mime_type="application/json"
            )
        )

        if not response or not response.text:
            raise Exception("AI response was empty or None")
        
        try:
            parsed = json.loads(response.text)
        except Exception as parse_err:
            raise Exception(f"Invalid JSON response from Gemini: {response.text}")

        return {
            "score": parsed.get("ats_score", 0),
            "keyword_match": parsed.get("keyword_match", 0),
            "missing_keywords": parsed.get("missing_keywords", []),
            "suggestions": parsed.get("suggestions", [])
        }

    except Exception as e:
        # Final fallback if AI fails entirely
        return {
            "score": 0,
            "keyword_match": 0,
            "missing_keywords": [],
            "suggestions": [f"Error analyzing CV: {str(e)}"]
        }

def extract_key_metrics(cv_content):
    """Extract quantifiable metrics from CV"""
    # Pattern to find numbers and percentages
    metrics_pattern = r'(\d+(?:\.\d+)?(?:%|K|M|B|k|m|b|\+|,\d+)*)'
    
    metrics = re.findall(metrics_pattern, cv_content)
    
    return {
        'total_metrics': len(metrics),
        'metrics_found': metrics,
        'quantification_score': min(100, len(metrics) * 5)  # 5 points per metric, max 100
    }

def enhance_action_verbs(content, intensity="High"):
    """Enhance action verbs in CV content"""
    
    action_verbs = {
        "Moderate": [
            "managed", "developed", "created", "implemented", "led", "coordinated",
            "designed", "analyzed", "improved", "organized", "planned", "supervised"
        ],
        "High": [
            "spearheaded", "orchestrated", "revolutionized", "transformed", "pioneered",
            "architected", "optimized", "streamlined", "accelerated", "amplified"
        ],
        "Very High": [
            "catapulted", "revolutionized", "masterminded", "propelled", "dominated",
            "commanded", "conquered", "devastated", "obliterated", "annihilated"
        ]
    }
    
    # This would be implemented with more sophisticated text processing
    # For now, return the content as-is
    return content

def generate_interview_qa(resume_text, job_description):
    """Generate interview Q&A using Gemini AI"""
    prompt = f"""
    You are an expert career coach and interviewer.

    TASK:
    Generate **exactly 20 interview questions and answers** for the candidate based on their resume and the job description.

    ✅ Structure:
    - 8 Behavioral questions (fitment, company, teamwork, problem-solving, adaptability)
    - 12 Technical questions based on ATS keywords, JD tools, frameworks, and skills.

    ✅ Format STRICTLY:
    Q1: [Behavioral Question]
    A1:
    - Point 1
    - Point 2
    - Point 3
    - Point 4
    - Point 5
    - Point 6

    Q2: [Next Question]
    A2:
    - ...

    ✅ Rules:
    - All answers MUST have **6 bullet points minimum**.
    - No repetition of questions or answers.
    - Technical questions should be advanced and role-specific.
    - Cover JD-specific tools, methods, and problem scenarios.
    - Include the most important ATS keywords in both questions and answers.

    Resume:
    {resume_text}

    Job Description:
    {job_description}
    """
    if not model:
        raise Exception("Gemini AI client not initialized")

    response = model.generate_content(
        prompt,  # or contents=prompt
        generation_config=types.GenerationConfig(
            temperature=0.2  # optional
        )
    )
    if not response or not response.text:
        raise Exception("AI response was empty")

    return response.text


def export_interview_qa(content):
    """Export Q&A content as PDF and DOCX"""
    from io import BytesIO
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
    from reportlab.lib.styles import getSampleStyleSheet
    from docx import Document

    # PDF Export
    pdf_buffer = BytesIO()
    doc = SimpleDocTemplate(pdf_buffer)
    styles = getSampleStyleSheet()
    story = []
    for line in content.split('\n'):
        if line.strip():
            story.append(Paragraph(line.strip(), styles['Normal']))
            story.append(Spacer(1, 12))
    doc.build(story)
    pdf_buffer.seek(0)

    # DOCX Export
    docx_buffer = BytesIO()
    word_doc = Document()
    for line in content.split('\n'):
        if line.strip():
            word_doc.add_paragraph(line.strip())
    word_doc.save(docx_buffer)
    docx_buffer.seek(0)

    return pdf_buffer, docx_buffer
